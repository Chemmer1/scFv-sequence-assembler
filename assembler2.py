import tkinter as tk #tkinter è una libreria per le interfacce GUI (graphical user interface)
from tkinter import filedialog, scrolledtext, messagebox
from io import StringIO #converte stringhe in file temporanei(?)

import pandas as pd #Per manipolare dati nelle tabelle
from Bio import SeqIO, pairwise2 #Sequence aligner deprecato, utile per ottenere seqA, seqB
from Bio.Align import PairwiseAligner #Sequence aligner di Biopython
from Bio.pairwise2 import format_alignment

import os #per interagire col sistema operativo e manipolare file (unisci, elimina, trova...)
import difflib #in uso? Serve per trovare la differenza tra stringhe
from datetime import datetime #Importa la data corrente

from docx import Document #Per generare file docx
from docx.shared import RGBColor, Pt, Inches #Testo colorato, Dimensioni testo, Intestamento pagina
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

from vquest.vq import vquest, layer_configs, DEFAULTS #libreria imgt/v per fare richieste

now = datetime.now() # Get current date and time
timestamp = now.strftime("%d-%m-%Y %H-%M") # Format it as string

def aligner(mode, match, miss, gap, xgap):
    aligner = PairwiseAligner()
    aligner.mode = mode
    aligner.match_score = match
    aligner.mismatch_score = miss
    aligner.open_gap_score = gap
    aligner.extend_gap_score = xgap
    return aligner  

def check_selection_data(): #Gestisce le variabili a seconda dei radiobuttons selezionati (per trimming e sorting)
    if radio_vector.get() == "dnl":
        expected_start = 106
        expected_end = 120
        first = "ggcgcgcatgcc".upper() #12bp del pDNL6,  
        second = "gctagcggcaaaccaatcccaaacc".upper() #25bp del pDNL6
        seq_limit = 46
    elif radio_vector.get() == "gex":
        expected_start = 82
        expected_end = 122    
        first = "ggcgcgcatgcc".upper() #12bp del pgex 
        second = "GCTAGCGATTACAAGGACGACGACGA".upper() #25bp del pgex
        seq_limit = 24             
    elif radio_vector.get() == "hygro":
        expected_start = 68
        expected_end = 60  
        first = "tctctccacaggtggcgcgc".upper() #20bp del phygro 
        second = "GCTAGCCCACCTAGTACATGTTCTA".upper() #25bp del phygro
        seq_limit = 40             
    elif radio_vector.get() == "custom":
        expected_start = len(input_up_boundary.get("1.0", tk.END).strip().upper())-1
        expected_end = len(input_down_boundary.get("1.0", tk.END).strip().upper())-1
        first = input_up_boundary.get("1.0", tk.END).strip()[-12:].upper()
        second = input_down_boundary.get("1.0", tk.END).strip()[:25].upper()
        seq_limit = 40    
    
    #linker radiobuttons----------------------------------------------------------------------------------    
    if radio_linker.get() == "custom":     
        linker_sequence = input_linker.get("1.0", tk.END).strip().upper()
        if len(linker_sequence) < 20:
            messagebox.showwarning("Warning", "Linker sequence is less than 20bp, alignment might be difficult due to aspecificity")
    elif radio_linker.get() == "opt1":      
        linker_sequence = "TCCGGAGGGTCGACCATAACTTCGTATAATGTATACTATACGAAGTTATCCTCGAGCGGTACC".upper()
    else:
        linker_sequence = ""
        messagebox.showerror("Error", "Error in linker sequence input")
    
    return expected_start, expected_end, first, second, seq_limit, linker_sequence

def trimmer():
    destinazione = filedialog.askopenfilenames(filetypes=[("ABI files", "*.ab1")]) #genera una tupla di files nelle directory
    trimmed_seq, file_name, trimmed_qual = [], [], []   
     
    # TRIM
    for path in destinazione:      
        filename = os.path.basename(path)
        record = SeqIO.read(path,"abi") #genera u   n oggetto analizzabile a partire dalla directory fornita
        sequenza = record.seq
        qualità = record.letter_annotations["phred_quality"]           
        expected_start, expected_end, *_ = check_selection_data()
        end_limit = len(sequenza) - expected_end + 20
        end = end_limit + 20
        start = expected_start - 40
        start_limit = expected_start -20

        #trimmer in base alla qualità
        while start < start_limit and qualità[start] < 50:
            start += 1
        while end > end_limit and qualità[end] < 50:   
            end -=1            
        file_name.append(filename)
        trimmed_seq.append(sequenza[start: end])
        trimmed_qual.append(qualità[start: end])
    if expected_start < 70:
        messagebox.showwarning("Alert", "By the given upstream sequence, the GOI sequence is expected to start before 70 nt, this might be problematic for the trimming start due to Sanger sequencing first base calls being noisy")
    if end_limit > 950:
        messagebox.showwarning("Alert", "The GOI sequence is expected to end > 950bp, this might be problematic for the trimming end due to Sanger's last base calls being noisy")
    
    return file_name, trimmed_seq, trimmed_qual   

def sorter(file_name, trimmed_seq, trimmed_qual):
    sense_seq, sense_name, sense_qual = [], [], []
    reverse_seq, reverse_name, reverse_qual = [], [], []

    for seq, name, qual in zip(trimmed_seq, file_name, trimmed_qual):
        try:
            *_, linker_sequence = check_selection_data()
            alignment = aligner("local", 1, -1, -5, -2).align(linker_sequence, seq)[0]
            score = alignment.score
            if score >= (len(linker_sequence)-3):
                sense_name.append(name)
                sense_seq.append(seq)
                sense_qual.append(qual)
            else:
                seq = seq.reverse_complement()
                reverse_name.append(name)
                reverse_seq.append(seq)
                qual = qual[::-1]
                reverse_qual.append(qual)
        except Exception as e:
            output_box.insert(tk.END, f"error {e}")
            return [],[],[],[]
    return sense_seq, reverse_seq, sense_name, reverse_name, sense_qual, reverse_qual

def finder(sense_seq, reverse_seq, sense_name, reverse_name, sense_qual, reverse_qual):  
    found_name, found_alignments, found_qual = [], [], []
    
    for Fseq, Fname, Fqual in zip(sense_seq, sense_name, sense_qual):
        best_pair = None
        best_score= float('-inf')
        best_Rname = None
        best_Rqual = None
        for Rseq, Rname, Rqual in zip(reverse_seq, reverse_name, reverse_qual):    
            score = aligner("global", 1, -2, -3, -2).score(Fseq, Rseq)
            if score > best_score:
                best_score = score
                best_pair = (Fseq, Rseq)
                best_Rname = Rname
                best_Rqual = Rqual
        if best_pair:
            found_name.append((Fname, best_Rname))
            found_alignments.append(best_pair) 
            found_qual.append((Fqual, best_Rqual))
            
    return found_alignments, found_name, found_qual

def merger(found_alignments, found_name, found_qual):       
    merged_name, merged_seq, merged_alignment = [], [], []
    
    for (Fseq, Rseq), (Fname, Rname), (Fqual, Rqual) in zip(found_alignments, found_name, found_qual):        
        alignment = pairwise2.align.globalms(str(Fseq), str(Rseq), 1, -2, -3, -2, one_alignment_only = True)[0]             
        aligned_F_seq = alignment.seqA
        aligned_R_seq = alignment.seqB       
        merged = []
        
        fi = ri = 0  # index for original quality lists
        for bf, br in zip(aligned_F_seq, aligned_R_seq):
            if bf != '-':
                qf = Fqual[fi]
                fi += 1
            else:
                qf = -1
            if br != '-':
                qr = Rqual[ri]
                ri += 1
            else:
                qr = -1
            # Choose base with better quality
            if qf >= qr:
                merged.append(bf)
            elif qf ==-1 and qr == -1:
                merged.append('N')
            else:
                merged.append(br)
        
        merged_alignment.append(alignment)
        merged_seq.append(''.join(merged))
        merged_name.append((Fname, Rname))

    return merged_name, merged_seq, merged_alignment
                           
def merged_trimmer(merged_name, merged_seq):    
    #INCLUDI PRIMER 46BP EXAMPLE
    VL_out, linker_out, VH_out = [], [], []
    *_, first, second, seq_limit, linker_sequence = check_selection_data()  
    seq_trim = []
    
    for(fname,rname), seq in zip(merged_name, merged_seq):
        alignment_first = aligner("local", 2, -1, -3, -1).align(first, seq[0:seq_limit])[0]     
        alignment_second = aligner("local", 1, -1, -5, -2).align(second, seq)[0]
        inizio = alignment_first.aligned[1][0][1]
        fine = alignment_second.aligned[1][-1][0]
        seq_trim.append(seq[inizio:fine]) 
    
    for seq, (fname, rname) in zip(seq_trim, merged_name):
        alignment = aligner("local", 1, -1, -5, -2).align(linker_sequence, seq)[0]
        inizio = alignment.aligned[1][0][0]
        fine = alignment.aligned[1][-1][1]
        VL_out.append(seq[0:inizio])
        linker_out.append(seq[inizio:fine])
        VH_out.append(seq[fine:])
    
    return VL_out, linker_out, VH_out, merged_name

def imgt_call(VL_out, VH_out, merged_name):    
    #Funzione che estrae i valori d'interesse da imgt/v
    def extract_value(tsv, colname):
        try:
            df = pd.read_csv(StringIO(tsv), sep="\t", low_memory=False) # stringIO converte in file la stringa tsv così che pandas (pd) può leggerlo e generare un dataframe. i dati sono separati da un Tab ed evita di leggere in chunks (legge tutto il dataframe)
            return df[colname].dropna().astype(str).str.split(",").str[0].iloc[0] #ottiene i dati dalla serie dataframe[chiamata] (imgt/v) puliti, dropna rimuove colonne vuote, convertite in stringhe e il primo valore 
        except Exception:
            return "N/A"
            
    VL_fasta, VH_fasta = [], []
    results_igh, results_igl = [], []
    HV, HJ, HCDR = [],[],[]
    LV, LJ, LCDR = [],[],[]
    
    #converto vl e vh in formato FASTA
    for vl,vh, (fname, rname) in zip(VL_out, VH_out, merged_name):
        vlf = f">VL\n{vl}\n"
        vhf = f">VL\n{vh}\n"
        VL_fasta.append(vlf)
        VH_fasta.append(vhf)
    
    #si generano dei dizionari da usare per sostituire i dati default di imgt/v   
    for vl, vh, (fname,rname) in zip(VL_fasta, VH_fasta, merged_name):     
        try:           
            #aggiungi alla lista la richiesta a imgt/v (vquest)
            results_igl.append(vquest(layer_configs(DEFAULTS,{"species": "human", "receptorOrLocusType": "IGL", "sequences": vl}))["vquest_airr.tsv"])
            results_igh.append(vquest(layer_configs(DEFAULTS,{"species": "human", "receptorOrLocusType": "IGH", "sequences": vh}))["vquest_airr.tsv"])
        except Exception as e:
            output_box.insert(tk.END, f"errore {e}")
            results_igh.append("")
            results_igl.append("")
    
    #estrae da imgt/v i valori d'interesse        
    for vl_tsv, vh_tsv, (fname, rname) in zip(results_igl, results_igh, merged_name):        
        HV.append(extract_value(vh_tsv, "v_call"))
        HJ.append(extract_value(vh_tsv, "j_call"))
        HCDR.append(extract_value(vh_tsv, "cdr3_aa"))
        LV.append(extract_value(vl_tsv, "v_call"))
        LJ.append(extract_value(vl_tsv, "j_call"))
        LCDR.append(extract_value(vl_tsv, "cdr3_aa"))
    
    return HV, HJ, HCDR, LV, LJ, LCDR
    
def save_to_word(VL_out, linker_out, VH_out, merged_name, merged_alignment, LV, LJ, LCDR, HV, HJ, HCDR):   
    doc = Document()   # Create a new Word document
    filename = f"pDNL6 Assembly {timestamp}.docx"
    
    #impostazione della pagina
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.5)   # Narrow left margin
        section.right_margin = Inches(0.5)  # Narrow right margin  
    
    #impostazione carattere
    styles = doc.styles
    if 'MyStyle' not in styles:
        style = styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Courier New'
        style.font.size = Pt(8)
    
    #titolo
    T = doc.add_paragraph(style = "MyStyle")
    T_run = T.add_run(f"pDNL6 ScFv Assembler (.ab1) - {timestamp}")
    T_run.font.size = Pt(12)
    T_run.bold = True
    doc.add_paragraph("")
     
    allineamento_output = []
    
    for (fseq,rseq), (fname,rname) in zip(merged_alignment, merged_name):
        alignment = aligner("global", 1, -2, -3, -2).align(fseq, rseq)[0]
        allineamento_output.append(alignment)
      
    for vl,lin,vh, (Fname, Rname), align, lv, lj, lcdr, hv, hj, hcdr in zip(VL_out, linker_out, VH_out, merged_name, allineamento_output, LV, LJ, LCDR, HV, HJ, HCDR):
        p = doc.add_paragraph(style = "MyStyle")
        p_run = p.add_run(f"File computati: {Fname} (SENSE) - {Rname} (ANTISENSE) \nConsensus sequence:\n")
        p_run.font.size = Pt(9)
        p_run.bold = True
        p_run2 = p.add_run(f"{vl}\n")
        p_run3 = p.add_run(f"{lin}\n")
        p_run3.font.color.rgb = RGBColor(0, 0, 255)
        p_run4 = p.add_run(f"{vh}\n\n")         
        p_run5 = p.add_run(f"Allineamento:\n{align}")
        
        #Tabella VL VH e geni
        table = doc.add_table(rows=3, cols=6)
        table.style = 'Table Grid'
        table.cell(0, 0).merge(table.cell(0, 2)).text = "VL"
        table.cell(0, 3).merge(table.cell(0, 5)).text = "VH"
        subheaders = ["V Gene", "J Gene", "CDR"] * 2
        for i, text in enumerate(subheaders):
            table.cell(1, i).text = text
        values = [lv, lj, lcdr, hv, hj, hcdr]
        for i, text in enumerate(values):
            table.cell(2, i).text = text
        doc.add_paragraph("")
    
    doc.save(filename) # Save the document
    
    return filename
    
def output(VL, linker_out, VH_out, merged_name, merged_alignment, filename):
    if VL and linker_out and VH_out and merged_name:
        path = os.path.abspath(filename)
        output_box.insert(tk.END, f"File {filename} creato con successo! \nSalvato in: {path}")
    else:
        output_box.insert(tk.END, f"C'è un errore, alcuni dati risultano vuoti (?)")
def check_input():
    return
    
def run_pipeline():
    if(radio_vector.get() == "custom" and (input_up_boundary.get("1.0", "end-1c") == placeholders[0] or input_down_boundary.get("1.0", "end-1c") == placeholders[1])) or (radio_linker.get() == "custom" and input_linker.get("1.0", "end-1c") == placeholders[2]):
        messagebox.showerror("Error", "You have to type in a nucleotide sequence in the custom textboxes!")
    else:
        file_name, trimmed_seq, trimmed_qual                                        = trimmer               ()
        sense_seq, reverse_seq, sense_name, reverse_name, sense_qual, reverse_qual  = sorter                (file_name, trimmed_seq, trimmed_qual)
        found_alignments, found_name, found_qual                                    = finder                (sense_seq, reverse_seq, sense_name, reverse_name, sense_qual, reverse_qual)
        merged_name, merged_seq, merged_alignment                                   = merger                (found_alignments, found_name, found_qual)
        VL_out, linker_out, VH_out, merged_name                                     = merged_trimmer        (merged_name, merged_seq)
        HV, HJ, HCDR, LV, LJ, LCDR                                                  = imgt_call             (VL_out, VH_out, merged_name)
        filename                                                                    = save_to_word          (VL_out, linker_out, VH_out, merged_name, found_alignments, LV, LJ, LCDR, HV, HJ, HCDR)
        output(VL_out, linker_out, VH_out, merged_name, merged_alignment, filename)

window = tk.Tk()
window.title("ScFv Assembler")

#----------------------------------------------------------GUI LOGIC organizzata in griglia, (non è possibile ordinare con .pack)----------------------------------------------------------------------------------------------------------
#------------------------LABELS------------------------------------------------------- 
tk.Label(window, font=("", 13), text="VECTOR BOUNDARIES SELECTION--------------------").grid(row = 0, column=0, sticky="w", padx=10, pady=10, columnspan = 3) #.grid dopo la definizione del widget lo rende None e non si può più modificarlo
tk.Label(window, font=("", 13), text="LINKER SEQUENCE INSERTION-------------------------").grid(row=0, column=1, sticky="w", padx=5, pady=5, columnspan = 3)

#------------------------- Radio button GUI -------------------------------------------------------
def check_selection_vector_GUI(textboxes):  
    #vector radiobuttons--------------------------------------------------------
    window.focus_set()   
    if radio_vector.get() == "custom":
        for i, box in enumerate(textboxes[:-1]):
            textboxes[i].config(state="normal")
            textboxes[i].delete("1.0", tk.END)
            placeholder_events(textboxes[i], placeholders[i])
    else:
        for i, box in enumerate(textboxes[:-1]):
            textboxes[i].delete("1.0", tk.END)
            textboxes[i].config(state="disabled")
            window.focus_set()   
        
    if (radio_vector.get() !="dummy") and (radio_linker.get() != "dummy"):
        run_button.config(state="normal")
    else:
        run_button.config(state="disabled")

def check_selection_linker_GUI():
    global linker_sequence
    window.focus_set()   
    #linker radiobuttons-------------------------------------------------------
    if radio_linker.get() == "custom":
        input_linker.config(state="normal")
        input_linker.delete("1.0", tk.END)
        placeholder_events(input_linker, placeholders[len(placeholders)-1])        
        linker_sequence = input_linker.get("1.0", tk.END).strip().upper()
    elif radio_linker.get() == "opt1":
        input_linker.delete("1.0", tk.END)
        input_linker.config(state="disabled")
        window.focus_set()        
        linker_sequence = "TCCGGAGGGTCGACCATAACTTCGTATAATGTATACTATACGAAGTTATCCTCGAGCGGTACC".upper()
    else:
        linker_sequence = ""
    
    if (radio_vector.get() !="dummy") and (radio_linker.get() != "dummy"):
        run_button.config(state="normal")
    else:
        run_button.config(state="disabled")
   
#Generazione dei bottoni radio per la scelta del vettore
radio_vector = tk.StringVar(value="dummy")
radio_vector_options = [("pDNL6 (automatic)","dnl"), ("pGex (automatic)","gex"), ("pHygro XbaI-Nhe (automatic)","hygro"), ("Custom boundaries","custom")]
for i, (text, value) in enumerate(radio_vector_options):
    rb = tk.Radiobutton(window, text=text, variable=radio_vector, value=value, command = lambda: check_selection_vector_GUI(textboxes))
    rb.grid(column = 0, row=i+1,padx=5, pady=5, sticky="w")

#Generazione dei bottoni radio per la scelta del linker
radio_linker = tk.StringVar(value = "dummy")
radio_linker_options = [("TCCGGAGGGTCGACCATAACTTCGTATAATGTATACTATACGAAGTTATCCTCGAGCGGTACC","opt1"), ("Custom linker","custom")]
for i, (text, value) in enumerate(radio_linker_options):
    rb = tk.Radiobutton(window, text=text, variable=radio_linker, value=value, command = lambda: check_selection_linker_GUI())
    rb.grid(column = 1, row=i+1,padx=5, pady=5, sticky="w")

#------------------------TEXTBOX------------------------------------------------------ 
def placeholder_events(text_widget, placeholder_text):
    def on_focus_in(event):
        if text_widget.get("1.0", "end-1c") == placeholder_text:
            text_widget.delete("1.0", tk.END)
            text_widget.config(fg="black")

    def on_focus_out(event):
        if text_widget.get("1.0", "end-1c").strip() == "":
            text_widget.insert("1.0", placeholder_text)
            text_widget.config(fg="gray")
    
    text_widget.bind("<FocusIn>", on_focus_in)
    text_widget.bind("<FocusOut>", on_focus_out)
    text_widget.insert("1.0", placeholder_text)
    text_widget.config(fg="gray")

#--------------------generazione delle textbox
placeholders =["Insert the whole expected upstream sequence from the GOI (ideally > 100nt)", "Insert the whole expected downstream sequence from the GOI", "Insert a custom linker sequence (ideally > 20nt)"]
textboxes = []
j=len(radio_vector_options)
for i, placeholder in enumerate(placeholders):
    box=scrolledtext.ScrolledText(window, width=70, height=1, state="disabled")
    if i == (len(placeholders)-1):
        box.grid(row= len(radio_linker_options)+1, column=1, padx=5, pady=5)
        placeholder_events(box, placeholder)
        textboxes.append(box)
    else:
        box.grid(row= j+1, column=0, padx=5, pady=5)
        placeholder_events(box, placeholder)
        textboxes.append(box)
        j += 1

global input_up_boundary
global input_down_boundary
input_up_boundary = textboxes[0]
input_down_boundary = textboxes[1]
input_linker = textboxes[2]

#------------------------- Output Box -------------------------
output_box = scrolledtext.ScrolledText(window, width=70, height=10, font=("Courier", 15))
output_box.grid(row=10, column=0, columnspan=6, padx=10, pady=10)

#------------------------- Run button -------------------------
run_button = tk.Button(window, text="Esegui!", width=40, height= 5, font=("Courier", 15), command = run_pipeline, state = "disabled")
run_button.grid(row = 5, column=1, rowspan=2, pady=10, sticky="w")


check_selection_vector_GUI(textboxes)
check_selection_linker_GUI()
window.mainloop()